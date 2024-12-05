from docx import Document

def generate_reader_card_file(reader, borrowed_books):
    """Создание Word-документа с информацией о карточке читателя."""
    # Создаем документ
    doc = Document()

    # Заголовок
    doc.add_heading("Карточка читателя", level=1)

    # Информация о читателе
    doc.add_paragraph(f"Имя: {reader.name}")
    doc.add_paragraph(f"Email: {reader.email}")
    doc.add_paragraph(f"Телефон: {reader.phone}")
    doc.add_paragraph(f"ID читателя: {reader.reader_id}")

    # История взятых книг
    doc.add_heading("История взятых книг", level=2)

    if borrowed_books:
        for record in borrowed_books:
            return_date = record.return_date if record.return_date else "Не возвращена"
            doc.add_paragraph(f"Книга: {record.book.title}")
            doc.add_paragraph(f"  Автор: {record.book.author}")
            doc.add_paragraph(f"  Жанр: {record.book.genre}")
            doc.add_paragraph(f"  Год: {record.book.year}")
            doc.add_paragraph(f"  Дата выдачи: {record.borrow_date}")
            doc.add_paragraph(f"  Дата возврата: {return_date}")
            doc.add_paragraph("-" * 40)
    else:
        doc.add_paragraph("У этого читателя нет взятых книг.")

    # Сохранение документа
    filename = f"Читатель_{reader.reader_id}_{reader.name.replace(' ', '_')}.docx"
    doc.save(filename)
    print(f"Файл карточки читателя сохранен как {filename}.")
