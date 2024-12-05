from datetime import datetime #для окончания срока карточки
from docx import Document #для создания карточки читателя в формате docx
import re
from models import Author, Book, BorrowRecord, Reader, Librarian


def add_author(authors):
    """Добавить автора."""
    name = input("Введите имя автора: ")
    birth_year = input("Введите год рождения автора: ")
    country = input("Введите страну автора: ")
    authors.append(Author(name, int(birth_year), country))
    print(f"Автор {name} добавлен.")


def remove_author(authors):
    """Удалить автора."""
    name = input("Введите имя автора для удаления: ")
    for author in authors:
        if author.name.lower() == name.lower():
            authors.remove(author)
            print(f"Автор {name} удален.")
            return
    print("Автор не найден.")


def add_book(books, authors):
    """Добавить книгу."""
    title = input("Введите название книги: ")
    author_name = input("Введите имя автора: ")
    genre = input("Введите жанр книги: ")
    year = input("Введите год издания книги: ")
    isbn = input("Введите ISBN книги: ")

    author = next((a for a in authors if a.name.lower() == author_name.lower()), None)
    if not author:
        print(f"Автор {author_name} не найден. Сначала добавьте автора.")
        return

    books.append(Book(title, author.name, genre, int(year), isbn))
    print(f"Книга '{title}' добавлена.")


def remove_book(books):
    """Удалить книгу."""
    title = input("Введите название книги для удаления: ")
    for book in books:
        if book.title.lower() == title.lower():
            books.remove(book)
            print(f"Книга '{title}' удалена.")
            return
    print("Книга не найдена.")


def register_reader(readers):
    """Зарегистрировать читателя."""
    name = input("Введите имя читателя: ")
    reader_id = len(readers) + 1
    email = input("Введите email читателя: ")
    phone = input("Введите телефон читателя: ")
    readers.append(Reader(name, reader_id, email, phone))
    print(f"Читатель {name} зарегистрирован.")

def show_reader_card(readers, borrow_records):
    """Просмотр карточки читателя по ID, включая историю взятых книг и создание файла."""
    reader_id = int(input("Введите ID читателя: "))  # Запрашиваем ID читателя

    # Ищем читателя по ID
    reader = next((r for r in readers if r.reader_id == reader_id), None)

    if reader:
        print("\nКарточка читателя:")
        print(f"  Имя: {reader.name}")
        print(f"  Email: {reader.email}")
        print(f"  Телефон: {reader.phone}")
        print(f"  ID читателя: {reader.reader_id}")

        # Срок действия карты (например, до 31.12.2027)
        card_expiry = datetime(2027, 12, 31)
        print(f"Срок действия карты: до {card_expiry.strftime('%d.%m.%Y')}")

        # Фильтруем все записи о книгах, выданных этому читателю
        borrowed_books = [record for record in borrow_records if record.reader == reader]

        if borrowed_books:
            print("\nИстория взятых книг:")
            for record in borrowed_books:
                return_date = record.return_date if record.return_date else "Не возвращена"
                print(f"  Книга: {record.book.title} — Автор: {record.book.author}, Жанр: {record.book.genre}, Год: {record.book.year}")
                print(f"  Дата выдачи: {record.borrow_date}")
                print(f"  Дата возврата: {return_date}")
                print("-" * 40)

            # Генерация и обновление файла карточки читателя
            generate_reader_card_file(reader, borrowed_books)
        else:
            print("\nУ этого читателя нет взятых книг.")
    else:
        print("Читатель с таким ID не найден.")


def lend_book(books, readers, librarians, borrow_records):
    """Выдача книги читателю."""
    book_title = input("Введите название книги: ")
    book = next((b for b in books if b.title == book_title), None)
    if not book:
        print("Книга не найдена.")
        return

    reader_id = input("Введите ID читателя: ")
    reader = next((r for r in readers if str(r.reader_id) == reader_id), None)
    if not reader:
        print("Читатель с таким ID не найден.")
        return

    librarian_id = input("Введите ID библиотекаря: ")
    librarian = next((l for l in librarians if str(l.librarian_id) == librarian_id), None)
    if not librarian:
        print("Библиотекарь с таким ID не найден.")
        return

    # Проверка, не выдана ли уже книга этому читателю
    for record in borrow_records:
        if record.book.title == book_title and record.reader.reader_id == reader.reader_id and not record.return_date:
            print(f"Ошибка: Книга '{book_title}' уже выдана читателю с ID {reader.reader_id}.")
            return

    borrow_date = input("Введите дату выдачи книги (ГГГГ-ММ-ДД): ")
    borrow_records.append(BorrowRecord(book, reader, librarian, borrow_date))
    print(f"Книга '{book_title}' успешно выдана читателю {reader.name}.")

def return_book(borrow_records):
    """Возврат книги."""
    book_title = input("Введите название возвращаемой книги: ")
    reader_id = input("Введите ID читателя: ")

    # Поиск записи о выдаче книги
    for record in borrow_records:
        if record.book.title == book_title and str(record.reader.reader_id) == reader_id and not record.return_date:
            return_date = input("Введите дату возврата книги (ГГГГ-ММ-ДД): ")
            record.return_date = return_date
            print(f"Книга '{book_title}' успешно возвращена читателем {record.reader.name}.")
            return

    print("Ошибка: Запись о выдаче книги не найдена или книга уже была возвращена.")

def show_all_books(books):
        """Вывести список всех книг в библиотеке."""
        if not books:
            print("В библиотеке пока нет книг.")
            return

        print("\nСписок всех книг в библиотеке:")
        for i, book in enumerate(books, start=1):
            print(f"{i}. {book.title} — Автор: {book.author}, Жанр: {book.genre}, Год: {book.year}, ISBN: {book.isbn}")

def show_all_data(books, authors, readers, librarians, borrow_records):
    """Вывести все данные."""
    print("\nАвторы:")
    for author in authors:
        print(f"  {author.name} ({author.birth_year}) - {author.country}")

    print("\nКниги:")
    for book in books:
        print(f"  {book.title} ({book.year}) - {book.author} [ISBN: {book.isbn}]")

    print("\nЧитатели:")
    for reader in readers:
        print(f"  {reader.name} (ID: {reader.reader_id}) - {reader.email}, {reader.phone}")

    print("\nБиблиотекари:")
    for librarian in librarians:
        print(f"  {librarian.name} (ID: {librarian.librarian_id}) - {librarian.email}, {librarian.phone}, {librarian.schedule}")

    print("\nЗаписи о выдаче:")
    for record in borrow_records:
        book_title = record.book.title if isinstance(record.book, Book) else record.book
        reader_name = record.reader.name if isinstance(record.reader, Reader) else record.reader
        librarian_name = record.librarian.name if isinstance(record.librarian, Librarian) else record.librarian
        status = f"Возвращена: {record.return_date}" if record.return_date else "На руках"
        print(f"  {book_title} -> {reader_name}, выдана: {librarian_name} ({status})")

def validate_email(email):
    #Проверка формата email.
    pattern = r"^[\w\.-]+@[\w\.-]+\.\w+$"
    return re.match(pattern, email)

def validate_phone(phone):
    # Проверка формата номера телефона (только цифры).
    return phone.isdigit()

def register_reader(readers):
    # Регистрация нового читателя.
    name = input("Введите имя читателя: ")
    reader_id = input("Введите ID читателя: ")
    email = input("Введите email читателя: ")
    while not validate_email(email):
        print("Неверный формат email. Попробуйте снова.")
        email = input("Введите email читателя: ")
    phone = input("Введите номер телефона читателя: ")
    while not validate_phone(phone):
        print("Номер телефона должен состоять только из цифр. Попробуйте снова.")
        phone = input("Введите номер телефона читателя: ")
    readers.append(Reader(name, int(reader_id), email, phone))
    print("Читатель успешно зарегистрирован.")

def generate_reader_card_file(reader, borrowed_books):
    """Создание Word-документа с информацией о карточке читателя."""
    # Создаем документ
    doc = Document()

    # Заголовок
    doc.add_heading("Карточка читателя", level=1)

    # Информация о читателе
    doc.add_paragraph(f"Имя: {reader.name}")
    doc.add_paragraph(f"Email: {reader.email}")
    doc.add_paragraph(f"Телефон: {reader.phone}")
    doc.add_paragraph(f"ID читателя: {reader.reader_id}")

    card_expiry = datetime(2027, 12, 31)
    doc.add_paragraph(f"Срок действия карты: до {card_expiry.strftime('%d.%m.%Y')}")

    # История взятых книг
    doc.add_heading("История взятых книг", level=2)

    if borrowed_books:
        for record in borrowed_books:
            return_date = record.return_date if record.return_date else "Не возвращена"
            doc.add_paragraph(f"Книга: {record.book.title}")
            doc.add_paragraph(f"  Автор: {record.book.author}")
            doc.add_paragraph(f"  Жанр: {record.book.genre}")
            doc.add_paragraph(f"  Год: {record.book.year}")
            doc.add_paragraph(f"  Дата выдачи: {record.borrow_date}")
            doc.add_paragraph(f"  Дата возврата: {return_date}")
            doc.add_paragraph("-" * 40)
    else:
        doc.add_paragraph("У этого читателя нет взятых книг.")

    # Сохранение документа
    filename = f"Читатель_{reader.reader_id}_{reader.name.replace(' ', '_')}.docx"
    doc.save(filename)
    print(f"Файл карточки читателя сохранен как {filename}.")


